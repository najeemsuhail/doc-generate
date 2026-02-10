import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os
import uuid

def create_customer_letters(excel_file, output_folder='output_letters'):
    """
    Read customer data from Excel and generate personalized Word documents.
    
    Expected Excel columns: SSA, Billing Account, CUSTOMER NAME, Accot Subtype, 
                          Department, Address, Status(Active/Inactive), 
                          Outstanding amount in Rs, CLOSURE DATE
    """
    
    # Create output folder if it doesn't exist
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    
    # Read Excel file
    try:
        df = pd.read_excel(excel_file)
    except FileNotFoundError:
        print(f"Error: {excel_file} not found!")
        return
    
    print(f"Found {len(df)} customers. Generating letters...\n")
    
    # Process each customer
    for idx, customer in df.iterrows():
        # Create a new Document
        doc = Document()
        
        # Add header (sender's details - can be customized)
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        header.add_run("[Your Company Name]\n[Your Address]\n[City, State, Pin]\n[Email/Phone]").font.size = Pt(10)
        
        # Add date
        doc.add_paragraph(f"\nDate: {datetime.now().strftime('%B %d, %Y')}\n")
        
        # Add recipient address
        recipient = doc.add_paragraph()
        recipient.alignment = WD_ALIGN_PARAGRAPH.LEFT
        customer_name = customer.get('CUSTOMER NAME', 'Valued Customer')
        recipient_text = f"{customer_name}\n"
        if pd.notna(customer.get('Address')):
            recipient_text += f"{customer['Address']}\n"
        recipient.add_run(recipient_text).font.size = Pt(11)
        
        # Add salutation
        salutation_name = str(customer_name).split()[0] if pd.notna(customer_name) else "Valued Customer"
        doc.add_paragraph(f"\nDear {salutation_name},")
        
        # Add body of letter based on status
        status = str(customer.get('Status(Active/Inactive)', 'Active')).lower().strip()
        outstanding = customer.get('Outstanding amount in Rs', 0)
        billing_account = customer.get('Billing Account', '')
        department = customer.get('Department', '')
        
        body_text = ""
        
        if 'inactive' in status:
            body_text = f"""We are writing to inform you that your account is currently inactive.

Account Details:
• Billing Account: {billing_account}
• Department: {department}
• Outstanding Amount: ₹{outstanding:,.2f}

If your account has been inactive due to closure or completion of services, please disregard this notice. However, if you have any outstanding payments, please settle them at your earliest convenience.

Payment can be made through:
• Bank transfer
• Check by mail
• Online payment portal
• Digital payment methods

If you have any questions regarding your account or need to reactivate your services, please contact us.

Thank you for your attention to this matter.
"""
        
        else:  # Active status
            body_text = f"""We are reaching out regarding your account status and outstanding balance.

Account Details:
• Billing Account: {billing_account}
• Department: {department}
• Outstanding Amount: ₹{outstanding:,.2f}
• Account Status: Active

Please review your account and ensure all payments are up to date. If you have any outstanding balance, we request you to settle it at your earliest convenience.

Payment Options:
• Bank transfer
• Check by mail
• Online payment portal
• Digital payment methods

If you have already made a payment or have any questions about your account, please feel free to contact us.

We value your business and look forward to a continued relationship with you.
"""
        
        doc.add_paragraph(body_text)
        
        # Add closing
        doc.add_paragraph(
            "Thank you for your prompt attention to this matter. We look forward to a continued relationship with you.\n\n"
            "Sincerely,\n\n"
            "[Your Name]\n"
            "[Your Title]\n"
            "[Company Name]"
        )
        
        # Save document with customer name and unique identifier
        customer_name = str(customer.get('CUSTOMER NAME', 'Customer')).replace(' ', '_').replace('/', '_')
        # Add billing account number and index to ensure uniqueness
        billing_account = str(customer.get('Billing Account', idx)).replace(' ', '_').replace('/', '_')
        filename = os.path.join(output_folder, f"Letter_{customer_name}_{billing_account}_{idx:03d}.docx")
        doc.save(filename)
        print(f"✓ Generated: {filename}")
    
    print(f"\n✓ All {len(df)} letters generated successfully in '{output_folder}' folder!")

if __name__ == "__main__":
    # Usage - update filename with your actual Excel file
    excel_file = 'your_file.xlsx'  # Change this to your actual file name
    create_customer_letters(excel_file, 'output_letters')
