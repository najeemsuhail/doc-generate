from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new Document
doc = Document()

# Add title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run("TEST TEMPLATE")
title_run.font.size = Pt(16)
title_run.font.bold = True

# Add company header
header = doc.add_paragraph()
header.alignment = WD_ALIGN_PARAGRAPH.CENTER
header_run = header.add_run("{COMPANY_NAME}")
header_run.font.size = Pt(12)

# Add date
doc.add_paragraph()
date_para = doc.add_paragraph("Date: {DATE}")
date_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

# Add recipient info
doc.add_paragraph()
recipient = doc.add_paragraph()
recipient.add_run("To:\n")
recipient.add_run("{CUSTOMER_NAME}\n")
recipient.add_run("{ADDRESS}")

# Add salutation
doc.add_paragraph()
salutation = doc.add_paragraph("Dear {CUSTOMER_NAME},")

# Add body
doc.add_paragraph()
body = doc.add_paragraph("""This is a test letter template to verify placeholder replacement works correctly.

Account Information:
• Billing Account: {BILLING_ACCOUNT}
• Department: {DEPARTMENT}
• Outstanding Amount: {OUTSTANDING_AMOUNT}
• Account Status: {STATUS}
• Closure Date: {CLOSURE_DATE}

This template contains all the standard placeholders. When you generate letters, each placeholder will be replaced with the corresponding customer data from your Excel file.

Please note that {CUSTOMER_NAME} should be replaced with the actual customer name.""")

# Add closing
doc.add_paragraph()
closing = doc.add_paragraph("""Sincerely,

{SENDER_NAME}
[Your Title]
{COMPANY_NAME}""")

# Save the template
doc.save('test_template.docx')
print("✓ Test template created: test_template.docx")
print("\nTemplate contains these placeholders:")
print("  {COMPANY_NAME}")
print("  {CUSTOMER_NAME}")
print("  {ADDRESS}")
print("  {BILLING_ACCOUNT}")
print("  {DEPARTMENT}")
print("  {OUTSTANDING_AMOUNT}")
print("  {STATUS}")
print("  {DATE}")
print("  {CLOSURE_DATE}")
print("  {SENDER_NAME}")
print("\nYou can now upload this template in the app to test placeholder replacement!")
