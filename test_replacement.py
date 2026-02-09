from docx import Document
from copy import deepcopy

# Load template and sample data
template_doc = Document('sample_template.docx')
print(f"Template has {len(template_doc.paragraphs)} paragraphs")

# Test replacement
replacements = {
    '{CUSTOMER NAME}': 'John Smith',
    '{COMPANY_NAME}': 'ABC Telecom',
    '{SENDER_NAME}': 'Billing Manager',
    '{Outstanding amount in Rs}': '5000.50',
    '{CLOSURE DATE}': '2026-03-01',
    '{Billing Account}': 'ACC-12345',
    '{Department}': 'Sales',
    '{Address}': '123 Main St, City',
    '{Status(Active/Inactive)}': 'Active',
    '{DATE}': '2026-02-09'
}

# Test on a copy
test_doc = deepcopy(template_doc)

print("\nBefore replacement:")
for para in test_doc.paragraphs[:5]:
    print(f"  {para.text}")

# Do replacement
def replace_text_in_paragraph(paragraph, replacements):
    full_text = paragraph.text
    needs_replacement = any(key in full_text for key in replacements.keys())
    
    if not needs_replacement:
        return False
    
    new_text = full_text
    for key, value in replacements.items():
        if key in new_text:
            new_text = new_text.replace(key, str(value))
            print(f"✓ Found and replaced: {key}")
    
    if new_text == full_text:
        return False
    
    for run in list(paragraph.runs):
        r = run._element
        r.getparent().remove(r)
    
    paragraph.add_run(new_text)
    return True

count = 0
for paragraph in test_doc.paragraphs:
    if replace_text_in_paragraph(paragraph, replacements):
        count += 1

for table in test_doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if replace_text_in_paragraph(paragraph, replacements):
                    count += 1

print(f"\nReplaced in {count} locations")

print("\nAfter replacement:")
for para in test_doc.paragraphs[:5]:
    print(f"  {para.text}")

test_doc.save('test_output.docx')
print("\n✅ Saved test_output.docx - Check if placeholders were replaced!")
