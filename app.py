import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os
import io
import zipfile
from pathlib import Path
from copy import deepcopy
import re

st.set_page_config(page_title="Customer Letter Generator", layout="wide", initial_sidebar_state="expanded")

# Custom CSS
st.markdown("""
    <style>
    .main {
        padding: 20px;
    }
    .stButton>button {
        width: 100%;
        height: 45px;
        border-radius: 5px;
        font-weight: bold;
    }
    </style>
    """, unsafe_allow_html=True)

st.title("📧 Customer Letter Generator")

# Helper function to replace text in Word document
def replace_text_in_paragraph(paragraph, replacements, debug=False):
    """Replace all placeholders in a paragraph, handling split runs"""
    # Get full paragraph text
    full_text = paragraph.text
    
    # Check if any replacement is needed
    needs_replacement = any(key in full_text for key in replacements.keys())
    
    if not needs_replacement:
        return False
    
    # Replace all placeholders in the full text
    new_text = full_text
    for key, value in replacements.items():
        if key in new_text:
            new_text = new_text.replace(key, str(value))
            if debug:
                st.write(f"  ✓ Replaced: {key} → {value}")
    
    # Only proceed if text actually changed
    if new_text == full_text:
        return False
    
    # Clear paragraph completely using XML manipulation
    for run in list(paragraph.runs):
        r = run._element
        r.getparent().remove(r)
    
    # Add the replaced text as a new run
    paragraph.add_run(new_text)
    return True

def replace_text_in_document(doc, replacements, debug=False):
    """Replace all placeholders in document with customer data"""
    replaced_count = 0
    
    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        if replace_text_in_paragraph(paragraph, replacements, debug):
            replaced_count += 1
    
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if replace_text_in_paragraph(paragraph, replacements, debug):
                        replaced_count += 1
    
    return replaced_count
    
    return doc

# Sidebar Navigation
with st.sidebar:
    st.header("📌 Navigation")
    menu = st.radio(
        "Select Feature:",
        ["🏠 Home", "📧 Generate Letters", "📚 Help"],
        label_visibility="collapsed"
    )
    st.divider()

# HOME PAGE
if menu == "🏠 Home":
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.header("Welcome to Customer Letter Generator!")
        st.markdown("""
        ### Quick Start Guide
        
        1. **Generate Letters** - Upload Excel file and create personalized Word documents
        2. **Settings** - Configure company details
        3. **Help** - View documentation and troubleshooting
        4. **About** - App information
        
        ### What This App Does
        - 📁 Upload customer Excel files
        - 📄 Use your own Word template OR text templates
        - 📄 Generate personalized Word documents
        - 📥 Download all letters as ZIP
        
        **Get Started:** Click "📧 Generate Letters" in the menu! ➜
        """)
    
    with col2:
        st.info("""
        **Features:**
        - ✅ Word template support
        - ✅ Text templates
        - ✅ Configurable dates
        - ✅ Bulk generation
        """)

# GENERATE LETTERS PAGE
elif menu == "📧 Generate Letters":
    st.markdown("Generate personalized Word documents for bulk mailing to customers")
    
    # Sidebar - Only date configuration
    with st.sidebar:
        st.subheader("📅 Letter Date")
        letter_date = st.date_input("Select date for letters", value=datetime.now().date())
        letter_date_str = letter_date.strftime('%B %d, %Y')

    # Step 1: Upload Excel
    st.header("📁 Step 1: Upload Excel File")
    uploaded_file = st.file_uploader("Choose your Excel file", type=['xlsx', 'xls'])
    
    if not uploaded_file:
        st.info("👆 Please upload an Excel file to get started")
        st.stop()
    
    df = pd.read_excel(uploaded_file)
    st.success(f"✓ File loaded successfully! ({len(df)} customers found)")
    
    with st.expander("📊 Preview Data", expanded=False):
        st.dataframe(df.head(10), use_container_width=True)
        st.info(f"Total rows: {len(df)}")

    # Step 2: Choose Template Source
    st.header("📋 Step 2: Upload Word Template")
    
    st.info("✓ Upload your own formatted Word document template with placeholders like {CUSTOMER_NAME}, {BILLING_ACCOUNT}, etc.")
    
    template_file = st.file_uploader(
        "Choose Word template file",
        type=['docx'],
        key="template_upload"
    )
    
    if not template_file:
        st.warning("Please upload a Word template file (.docx)")
        st.stop()
    
    try:
        # Save uploaded file to temporary location (Streamlit requirement)
        import tempfile
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            tmp.write(template_file.getbuffer())
            temp_template_path = tmp.name
        
        template_doc = Document(temp_template_path)
        st.success("✓ Template loaded successfully!")
        
        # Extract placeholders from template
        placeholders_found = set()
        for paragraph in template_doc.paragraphs:
            matches = re.findall(r'\{[^}]+\}', paragraph.text)
            placeholders_found.update(matches)
        
        for table in template_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        matches = re.findall(r'\{[^}]+\}', paragraph.text)
                        placeholders_found.update(matches)
        
        if placeholders_found:
            available_placeholders = sorted(list(placeholders_found))
            st.info(f"Found placeholders: {', '.join(available_placeholders)}")
        else:
            st.warning("No placeholders found in template. Use format: {PLACEHOLDER_NAME}")
                
    except Exception as e:
        st.error(f"Error loading template: {str(e)}")
        st.stop()

    col1, col2, col3 = st.columns(3)
    
    with col1:
        start_row = st.number_input("Start from row", min_value=1, max_value=len(df), value=1)
    
    with col2:
        end_row = st.number_input("End at row", min_value=1, max_value=len(df), value=len(df))
    
    with col3:
        st.empty()
    
    if st.button("🎯 Generate Letters", key="generate_btn"):
        progress_bar = st.progress(0)
        status_text = st.empty()
        generated_files = []
        
        try:
            for idx, (_, customer) in enumerate(df.iloc[start_row-1:end_row].iterrows()):
                progress = (idx + 1) / (end_row - start_row + 1)
                progress_bar.progress(progress)
                status_text.text(f"Generating letter {idx + 1} of {end_row - start_row + 1}...")
                
                customer_name = customer.get('CUSTOMER NAME', 'Valued Customer')
                outstanding = customer.get('Outstanding amount in Rs', 0)
                billing_account = customer.get('Billing Account', '')
                department = customer.get('Department', '')
                address = customer.get('Address', '')
                landline = customer.get('Landline', '')
                status = str(customer.get('Status(Active/Inactive)', 'Active')).lower().strip()
                
                # Create replacement dictionary - automatically from all Excel columns
                replacements = {}
                
                # Add all columns as placeholders (with various formatting)
                for col_name in df.columns:
                    col_value = str(customer.get(col_name, ''))
                    
                    # Create placeholders with different formats
                    replacements[f'{{{col_name}}}'] = col_value  # {CUSTOMER NAME}
                    replacements[f'{{{col_name.replace(" ", "_")}}}'] = col_value  # {CUSTOMER_NAME}
                    replacements[f'{{{col_name.replace(" ", "")}}}'] = col_value  # {CUSTOMERNAME}
                    replacements[f'{{{col_name.upper()}}}'] = col_value  # Uppercase
                    replacements[f'{{{col_name.lower()}}}'] = col_value  # Lowercase
                
                # Also add special formatted versions for outstanding amount
                try:
                    outstanding_amount = float(customer.get('Outstanding amount in Rs', 0))
                    replacements['{Outstanding amount in Rs}'] = f"{outstanding_amount:,.2f}"
                    replacements['{outstanding:,.2f}'] = f"{outstanding_amount:,.2f}"
                    replacements['{outstanding}'] = f"{outstanding_amount:,.2f}"
                except:
                    pass
                
                # Add landline field variations
                replacements['{Landline}'] = landline
                replacements['{LANDLINE}'] = landline
                replacements['{landline}'] = landline
                
                # Add date
                replacements['{DATE}'] = letter_date_str
                replacements['{date}'] = letter_date_str
                
                # Load fresh template for each customer (instead of deepcopy)
                doc = Document(temp_template_path)
                replaced_count = replace_text_in_document(doc, replacements, debug=False)
                
                # Save document
                filename = f"Letter_{str(customer_name).replace(' ', '_').replace('/', '_')}.docx"
                doc.save(filename)
                generated_files.append(filename)
            
            status_text.success(f"✅ Generated {len(generated_files)} letters successfully!")
            progress_bar.empty()
            
            if generated_files:
                zip_buffer = io.BytesIO()
                with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                    for file in generated_files:
                        zip_file.write(file, arcname=file)
                
                zip_buffer.seek(0)
                
                st.download_button(
                    label="📥 Download All Letters (ZIP)",
                    data=zip_buffer.getvalue(),
                    file_name=f"customer_letters_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip",
                    mime="application/zip",
                    key="download_zip"
                )
                
                for file in generated_files:
                    if os.path.exists(file):
                        os.remove(file)
        
        except Exception as e:
            st.error(f"❌ Error: {str(e)}")



# HELP PAGE
elif menu == "📚 Help":
    st.header("📚 Help & Documentation")
    
    with st.expander("📁 How to Prepare Excel File"):
        st.markdown("""
        Your Excel file should have these columns:
        
        | Column | Type | Example |
        |--------|------|---------|
        | SSA | Text | SSA-001 |
        | Billing Account | Text | ACC-12345 |
        | CUSTOMER NAME | Text | John Smith |
        | Accot Subtype | Text | Premium |
        | Department | Text | Sales |
        | Address | Text | 123 Main St |
        | Landline | Text | 040-12345678 |
        | Status(Active/Inactive) | Text | Active |
        | Outstanding amount in Rs | Number | 5000.00 |
        | CLOSURE DATE | Date | 2026-03-01 |
        """)
    
    with st.expander("📄 How to Create Word Template"):
        st.markdown("""
        1. **Open Microsoft Word** and create your letter template
        
        2. **Add placeholders** using column names from your Excel:
           - Use format: `{COLUMN_NAME}`
           - Can use spaces: `{CUSTOMER NAME}`
           - Or underscores: `{CUSTOMER_NAME}`
           - Or no spaces: `{CUSTOMERNAME}`
           - Case insensitive - all work!
        
        3. **Common placeholders:**
           - `{CUSTOMER NAME}` - Customer name
           - `{Billing Account}` - Account number
           - `{Outstanding amount in Rs}` - Outstanding amount
           - `{Department}` - Department
           - `{Address}` - Customer address
           - `{Landline}` - Landline number
           - `{CLOSURE DATE}` - Closure date
           - `{Status(Active/Inactive)}` - Account status
           - **Any column from your Excel file!**
        
        4. **Example template:**
           ```
           Dear {CUSTOMER NAME},
           
           Outstanding amount: Rs. {Outstanding amount in Rs}/-
           Account: {Billing Account}
           Department: {Department}
           
           Regards
           ```
        
        5. **Save as .docx file** (Microsoft Word format)
        6. **Upload in app** - it will auto-detect all placeholders!
        """)
    
    with st.expander("✏️ How to Use Text Templates"):
        st.markdown("""
        1. Go to **📧 Generate Letters**
        2. Select **📝 Use Text Template**
        3. Choose from 4 pre-built templates:
           - Payment Reminder
           - Collection Notice
           - Account Status
           - Service Closure
        4. Customize the text if needed
        5. Generate letters
        """)
    
    with st.expander("📥 Download & Print"):
        st.markdown("""
        1. Generate letters - downloaded as ZIP
        2. Extract the ZIP file
        3. Open each Word document
        4. Customize if needed
        5. Print for mailing
        """)
    
    with st.expander("❓ Troubleshooting"):
        st.markdown("""
        **Q: Column names don't match?**
        A: Make sure Excel columns match exactly (case-sensitive)
        
        **Q: Placeholders not being replaced?**
        A: Check spelling of placeholders and use curly braces {PLACEHOLDER}
        
        **Q: Can't upload Word template?**
        A: Make sure it's a .docx file (not .doc or .pdf)
        """)

# Footer
st.divider()
st.markdown("""
<div style='text-align: center; color: gray; font-size: 12px;'>
    Customer Letter Generator v2.0 | Word Template Edition
</div>
""", unsafe_allow_html=True)
