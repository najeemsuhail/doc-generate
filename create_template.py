from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new Document
doc = Document()

# Add title
title = doc.add_heading('PAYMENT REMINDER NOTICE', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add company header
company = doc.add_paragraph()
company.add_run('Company: {COMPANY_NAME}\n').bold = True
company.add_run('Contact: {SENDER_NAME}').italic = True

doc.add_paragraph()  # Blank line

# Add to address
doc.add_paragraph('To,')
doc.add_paragraph('{CUSTOMER NAME}')
doc.add_paragraph('{Address}')
doc.add_paragraph('{Department}')

doc.add_paragraph()  # Blank line

# Add date
date_para = doc.add_paragraph()
date_para.add_run('Date: ').bold = True
date_para.add_run('{DATE}')

doc.add_paragraph()  # Blank line

# Subject
subject = doc.add_paragraph()
subject.add_run('Subject: ').bold = True
subject.add_run('Outstanding Payment Reminder for Landline Number')

doc.add_paragraph()  # Blank line

# Body
doc.add_paragraph('Dear Sir/Madam,')

body = doc.add_paragraph()
body.add_run('This is to remind you that there is an outstanding balance against your account. Below are the details:')

doc.add_paragraph()  # Blank line

# Add table for details
table = doc.add_table(rows=6, cols=2)
table.style = 'Light Grid Accent 1'

# Table rows
rows = table.rows
rows[0].cells[0].text = 'Billing Account'
rows[0].cells[1].text = '{Billing Account}'

rows[1].cells[0].text = 'Customer Name'
rows[1].cells[1].text = '{CUSTOMER NAME}'

rows[2].cells[0].text = 'Outstanding Amount'
rows[2].cells[1].text = 'Rs. {Outstanding amount in Rs}/-'

rows[3].cells[0].text = 'Department'
rows[3].cells[1].text = '{Department}'

rows[4].cells[0].text = 'Account Status'
rows[4].cells[1].text = '{Status(Active/Inactive)}'

rows[5].cells[0].text = 'Service Closure Date'
rows[5].cells[1].text = '{CLOSURE DATE}'

doc.add_paragraph()  # Blank line

# Closing
closing = doc.add_paragraph('Please settle the outstanding amount at your earliest convenience to avoid service disruption.')
closing_2 = doc.add_paragraph('For any queries, please contact us.')

doc.add_paragraph()  # Blank line

# Signature
doc.add_paragraph('Thank you for your business.')
doc.add_paragraph()
sig = doc.add_paragraph('Regards,')
doc.add_paragraph('{SENDER_NAME}')

# Save
doc.save('sample_template.docx')
print("✅ Sample template created: sample_template.docx")
print("\nPlaceholders in template:")
print("  {COMPANY_NAME}")
print("  {SENDER_NAME}")
print("  {DATE}")
print("  {CUSTOMER NAME}")
print("  {Address}")
print("  {Department}")
print("  {Billing Account}")
print("  {Outstanding amount in Rs}")
print("  {Status(Active/Inactive)}")
print("  {CLOSURE DATE}")
